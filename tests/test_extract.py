"""Tests for text extraction pipeline — stub detection, full_text building, and format extraction."""

import io
import sqlite3
from unittest.mock import patch

import pytest

from regscope.db import get_db
from regscope.pipeline.extract import (
    _build_full_text,
    _extract_docx,
    _extract_html,
    _extract_pdf,
    _is_stub_text,
)


@pytest.fixture
def test_db(tmp_path):
    """Create a test database with a docket."""
    config = {"data": {"data_dir": str(tmp_path)}}
    db = get_db("TEST-EXTRACT", config)
    db.execute("INSERT INTO dockets (docket_id, title) VALUES ('TEST-EXTRACT', 'Extract Test')")
    db.commit()
    return db


class TestStubTextDetection:
    """Test that 'see attached' stub text is detected and replaced."""

    def test_see_attached_is_stub(self) -> None:
        """'See attached' variants should be recognized as stubs."""
        assert _is_stub_text("See attached.")
        assert _is_stub_text("See attached file")
        assert _is_stub_text("Please see the attached document")
        assert _is_stub_text("Attached is my comment")
        assert _is_stub_text("Comment attached")

    def test_real_comment_not_stub(self) -> None:
        """Real substantive comments should not be flagged as stubs."""
        assert not _is_stub_text("I support this rule because it protects public health.")
        assert not _is_stub_text("")  # empty is not a stub
        assert not _is_stub_text("A" * 300)  # long text is not a stub

    def test_stub_replaced_by_attachment_text(self, test_db) -> None:
        """When comment body is a stub and attachment text exists, full_text
        should contain only the attachment text, not the stub."""
        # Insert a comment with a stub body
        test_db.execute(
            """INSERT INTO comments
               (comment_id, docket_id, comment_text, detail_fetched)
               VALUES ('STUB-1', 'TEST-EXTRACT', 'See attached.', 1)"""
        )
        # Insert an attachment with extracted text
        test_db.execute(
            """INSERT INTO attachments
               (comment_id, file_url, file_format, extracted_text)
               VALUES ('STUB-1', 'http://example.com/doc.pdf', 'pdf',
                       'This is the real substantive comment from the PDF attachment.')"""
        )
        test_db.commit()

        _build_full_text(test_db, "TEST-EXTRACT")

        row = test_db.execute(
            "SELECT full_text FROM comments WHERE comment_id = 'STUB-1'"
        ).fetchone()
        full_text = row[0]

        # The stub text should NOT appear in full_text
        assert "see attached" not in full_text.lower()
        # The attachment text SHOULD appear
        assert "real substantive comment from the PDF" in full_text


class TestPdfExtraction:
    """Test PDF text extraction with PyMuPDF."""

    def test_extract_text_from_pdf(self) -> None:
        """Text inserted into a PDF via PyMuPDF should be extractable."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test document content for extraction.")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = _extract_pdf(pdf_bytes)
        assert "Test document content for extraction" in result

    def test_extract_multipage_pdf(self) -> None:
        """Text from multiple pages should be concatenated."""
        import fitz

        doc = fitz.open()
        page1 = doc.new_page()
        page1.insert_text((72, 72), "First page content.")
        page2 = doc.new_page()
        page2.insert_text((72, 72), "Second page content.")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = _extract_pdf(pdf_bytes)
        assert "First page content" in result
        assert "Second page content" in result

    def test_scanned_pdf_triggers_ocr_fallback(self) -> None:
        """A PDF with very little extractable text should attempt OCR."""
        import fitz

        # Create a PDF with no text (just a blank page) — simulates a scanned doc
        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        # Mock _ocr_pdf so we don't need tesseract installed
        with patch("regscope.pipeline.extract._ocr_pdf", return_value="OCR extracted text") as mock_ocr:
            result = _extract_pdf(pdf_bytes)
            mock_ocr.assert_called_once_with(pdf_bytes)
            assert result == "OCR extracted text"

    def test_scanned_pdf_with_minimal_text(self) -> None:
        """A PDF with fewer than 50 chars of extractable text should trigger OCR."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        # Insert a very short string (< 50 chars after stripping)
        page.insert_text((72, 72), "Hi")
        pdf_bytes = doc.tobytes()
        doc.close()

        with patch("regscope.pipeline.extract._ocr_pdf", return_value="Full OCR text") as mock_ocr:
            result = _extract_pdf(pdf_bytes)
            mock_ocr.assert_called_once()
            assert result == "Full OCR text"

    def test_pdf_extraction_error_returns_empty(self) -> None:
        """Invalid PDF content should return an empty string, not raise."""
        result = _extract_pdf(b"not a valid pdf")
        assert result == ""


class TestDocxExtraction:
    """Test DOCX text extraction with python-docx."""

    def test_extract_paragraphs(self) -> None:
        """Paragraphs in a DOCX file should be extracted."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = _extract_docx(docx_bytes)
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_extract_table_content(self) -> None:
        """Table cell contents should be extracted from DOCX."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Header paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = _extract_docx(docx_bytes)
        assert "Header paragraph" in result
        # Table cells are joined with " | " per row
        assert "A" in result
        assert "B" in result
        assert "C" in result
        assert "D" in result

    def test_extract_empty_docx(self) -> None:
        """An empty DOCX should produce an empty string."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = _extract_docx(docx_bytes)
        assert result == ""

    def test_docx_extraction_error_returns_empty(self) -> None:
        """Invalid DOCX content should return an empty string, not raise."""
        result = _extract_docx(b"not a valid docx file")
        assert result == ""


class TestHtmlExtraction:
    """Test HTML text extraction with script/style removal."""

    def test_extract_plain_html(self) -> None:
        """Basic HTML tags should be stripped, text content preserved."""
        html = b"<html><body><h1>Title</h1><p>Some content here.</p></body></html>"
        result = _extract_html(html)
        assert "Title" in result
        assert "Some content here" in result
        assert "<h1>" not in result
        assert "<p>" not in result

    def test_script_tags_removed(self) -> None:
        """Script tag contents should be completely removed."""
        html = b"""<html><body>
            <p>Visible text</p>
            <script type="text/javascript">
                var secret = "this should not appear";
                alert("hidden");
            </script>
            <p>More visible text</p>
        </body></html>"""
        result = _extract_html(html)
        assert "Visible text" in result
        assert "More visible text" in result
        assert "secret" not in result
        assert "alert" not in result
        assert "<script" not in result

    def test_style_tags_removed(self) -> None:
        """Style tag contents should be completely removed."""
        html = b"""<html><head>
            <style>
                body { color: red; font-size: 14px; }
                .hidden { display: none; }
            </style>
        </head><body><p>Styled content</p></body></html>"""
        result = _extract_html(html)
        assert "Styled content" in result
        assert "color: red" not in result
        assert "display: none" not in result
        assert "<style" not in result

    def test_whitespace_normalization(self) -> None:
        """Multiple whitespace characters should be collapsed."""
        html = b"<div>  word1   word2   word3  </div>"
        result = _extract_html(html)
        assert "word1 word2 word3" in result

    def test_nested_html_tags(self) -> None:
        """Deeply nested tags should be fully stripped."""
        html = b"<div><span><em><strong>Deep content</strong></em></span></div>"
        result = _extract_html(html)
        assert "Deep content" in result
        assert "<" not in result

    def test_mixed_script_and_style_with_content(self) -> None:
        """Both script and style tags should be removed while keeping body text."""
        html = b"""<html>
            <head><style>.x{color:blue}</style></head>
            <body>
                <script>var x=1;</script>
                <p>Keep this text</p>
                <style>.y{font:bold}</style>
                <script>console.log('hi');</script>
                <p>And this text</p>
            </body>
        </html>"""
        result = _extract_html(html)
        assert "Keep this text" in result
        assert "And this text" in result
        assert "var x" not in result
        assert "color:blue" not in result
        assert "console.log" not in result


class TestBuildFullText:
    """Test the full_text assembly logic beyond stub replacement."""

    def test_normal_comment_without_attachment(self, test_db) -> None:
        """A normal comment body without attachments should become full_text."""
        test_db.execute(
            """INSERT INTO comments
               (comment_id, docket_id, comment_text, detail_fetched)
               VALUES ('NORMAL-1', 'TEST-EXTRACT',
                       'I support this proposed regulation for environmental reasons.', 1)"""
        )
        test_db.commit()

        _build_full_text(test_db, "TEST-EXTRACT")

        row = test_db.execute(
            "SELECT full_text FROM comments WHERE comment_id = 'NORMAL-1'"
        ).fetchone()
        assert "support this proposed regulation" in row[0]

    def test_comment_with_attachment_concatenates(self, test_db) -> None:
        """A non-stub comment with attachments should concatenate body + attachment text."""
        test_db.execute(
            """INSERT INTO comments
               (comment_id, docket_id, comment_text, detail_fetched)
               VALUES ('CONCAT-1', 'TEST-EXTRACT',
                       'Please consider these detailed comments on the proposed rule.', 1)"""
        )
        test_db.execute(
            """INSERT INTO attachments
               (comment_id, file_url, file_format, extracted_text)
               VALUES ('CONCAT-1', 'http://example.com/detailed.pdf', 'pdf',
                       'Additional detailed analysis from the attached document.')"""
        )
        test_db.commit()

        _build_full_text(test_db, "TEST-EXTRACT")

        row = test_db.execute(
            "SELECT full_text FROM comments WHERE comment_id = 'CONCAT-1'"
        ).fetchone()
        full_text = row[0]

        # Both body and attachment text should appear
        assert "consider these detailed comments" in full_text
        assert "Additional detailed analysis" in full_text

    def test_stub_without_attachment_gets_empty_text(self, test_db) -> None:
        """A stub comment with no extracted attachment text should get empty full_text."""
        test_db.execute(
            """INSERT INTO comments
               (comment_id, docket_id, comment_text, detail_fetched)
               VALUES ('STUB-NOATT-1', 'TEST-EXTRACT', 'See attached.', 1)"""
        )
        test_db.commit()

        _build_full_text(test_db, "TEST-EXTRACT")

        row = test_db.execute(
            "SELECT full_text FROM comments WHERE comment_id = 'STUB-NOATT-1'"
        ).fetchone()
        assert row[0] == ""

    def test_comment_with_null_text(self, test_db) -> None:
        """A comment with NULL comment_text should not crash, and should use attachment if available."""
        test_db.execute(
            """INSERT INTO comments
               (comment_id, docket_id, comment_text, detail_fetched)
               VALUES ('NULL-1', 'TEST-EXTRACT', NULL, 1)"""
        )
        test_db.execute(
            """INSERT INTO attachments
               (comment_id, file_url, file_format, extracted_text)
               VALUES ('NULL-1', 'http://example.com/doc.pdf', 'pdf',
                       'Content from the attachment when body is null.')"""
        )
        test_db.commit()

        _build_full_text(test_db, "TEST-EXTRACT")

        row = test_db.execute(
            "SELECT full_text FROM comments WHERE comment_id = 'NULL-1'"
        ).fetchone()
        assert "Content from the attachment" in row[0]
